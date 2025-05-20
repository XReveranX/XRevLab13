import pyforms
from pyforms          import BaseWidget
from pyforms.controls import ControlText
from pyforms.controls import ControlButton
from pyforms import start_app
from docx import Document
from abc import ABC, abstractmethod
from pokets.Iron import Ironf
from pokets.TV import TVf
from pokets.WM import WNf




#self._iron_result.value = str(Ironf(int(self._iron.value),(int(self._tarif.value))))
#        self._tv_result.value = str(TVf(int(self._tv.value),(int(self._tarif.value))))
#        self._wm_result.value = str(WNf(int(self._wm.value),(int(self._tarif.value))))





doc = Document()

class programm(BaseWidget):

    def __init__(self):
        super(programm,self).__init__('programm')
        #Definition of the forms fields
        self._iron = ControlText('Введите время использования утюга (ч)')
        self._tv = ControlText('Введите время использования телевизора (ч)')
        self._wm = ControlText('Введите время использования стиральной машины (ч)')
        self._tarif = ControlText('Введите ваш тариф (Руб/кВт*ч)')
        self._button_rashit = ControlButton('Рассчитать')
        self._button_cancel = ControlButton('Очистить экран')
        self._iron_result = ControlText('Затраты на использование утюга (руб)','0')
        self._tv_result = ControlText('Затраты на использование телевизора (руб)','0')
        self._wm_result = ControlText('Затраты на использование стиральной машины (руб)','0')
        self._button_docx = ControlButton('сделать отчёт docx')
        self._ironr = self._iron.value

        self._button_rashit.value = self.__BA_rashit
        self._button_cancel.value = self.__BA_cancel
        self._button_docx.value = self.__BA_docx

    def __BA_cancel(self):
        """Button 'cancel' action event"""
        programm.close(self)

    def __BA_rashit(self):
        ironr._calc_res(self)

    def __BA_docx(self):
        """Button 'docx' action event"""
        doc.add_paragraph(("Затраты на использование утюга ", str(Ironf(int(self._iron.value),(int(self._tarif.value))))))
        doc.add_paragraph(("Затраты на использование телевизора ", str(TVf(int(self._tv.value),(int(self._tarif.value))))))
        doc.add_paragraph(("Затраты на использование стиральной машины ", str(WNf(int(self._wm.value),(int(self._tarif.value))))))
        doc.save("otchet.docx")

    @abstractmethod
    def calc_res(self):
        pass

class rashit(programm):
    def __init__(self):
        super().__init__()

    def _calc_res(self):
        pass


class ironr(rashit):
    def __init__(self):
        super().__init__()

    def calc_res(self):




#Execute the application
if __name__ == "__main__":   
    pyforms.start_app(programm, geometry=(200, 200, 600, 600))
